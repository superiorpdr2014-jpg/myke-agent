"""
Update SUPERIOR_ECHO_Prompt_0529.docx
新增 Step 7（包膜詢問）+ Step 8（板金烤漆歷史 + 等待真人評估）
插入位置：第六步（照片）之後、第七步（AI 分析）之前
"""

import sys, io, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement

SRC  = r"C:\Users\User\Desktop\ECHO AI建置\SUPERIOR_ECHO_Prompt_0529.docx"
DEST = r"C:\Users\User\Downloads\Myke_Agent\SUPERIOR_ECHO_Prompt_0529.docx"

doc = Document(SRC)

def para_text(p):
    return "".join(r.text for r in p.runs)

def make_para(text, bold=False):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    return p

def insert_before(ref_para, lines):
    parent = ref_para._p.getparent()
    idx = list(parent).index(ref_para._p)
    bold_set = {0}  # which line indices to bold
    for i, line in enumerate(lines):
        bold = i in bold_set and line.startswith("【")
        parent.insert(idx + i, make_para(line, bold=bold))

# ── locate anchor：第七步 AI 分析（剛才新增的）──────────────────────────────
anchor_ai = None
for p in doc.paragraphs:
    if "第七步：AI 照片分析" in para_text(p):
        anchor_ai = p
        break

print(f"anchor_ai found: {anchor_ai is not None}")

# ── 新增內容 ────────────────────────────────────────────────────────────────

NEW_STEPS = [
    "━━━━━━━━━━",
    "【第七步：詢問愛車是否有包膜或犀牛皮】",
    "━━━━━━━━━━",
    "",
    "「請問您的愛車目前是否有包覆包膜（PPF）或犀牛皮呢？」",
    "",
    "→ 若客戶回答「有」：",
    "  ECHO 須提醒：「了解，包膜狀況會影響修復方式，我們會一併列入評估，謝謝。」",
    "  並記錄：愛車有包膜 / 犀牛皮",
    "",
    "→ 若客戶回答「沒有」：",
    "  記錄：無包膜，繼續下一步",
    "",
    "→ 若客戶不確定：",
    "  ECHO：「沒關係，到店後技師會確認。」，繼續下一步",
    "",
    "━━━━━━━━━━",
    "【第八步：詢問是否曾進行過板金烤漆維修】",
    "━━━━━━━━━━",
    "",
    "「請問這台車以往有沒有做過板金烤漆維修的記錄呢？」",
    "",
    "→ 若客戶回答「有」：",
    "  ECHO：「了解，有過板金烤漆記錄我們也會一起評估，沒問題。」",
    "  記錄：有板金烤漆歷史",
    "",
    "→ 若客戶回答「沒有」或「不確定」：",
    "  記錄後繼續",
    "",
    "收集完成後，ECHO 回覆：",
    "「感謝您提供這些資訊，您愛車的狀況我們都已記錄起來了。",
    "技師收到資料後會盡快與您聯繫，評估最適合的修復方案，請稍候。」",
    "",
    "硬性規定：",
    "- 第七、八步為資訊收集步驟，不影響派單（[SEND_TO_JAY] 在三項必要資料齊全後已觸發）",
    "- 收集完成後 ECHO 等待真人技師進來評估，不再主動繼續詢問",
    "- 若三項必要資料尚未觸發派單，完成第八步後須立即觸發 [SEND_TO_JAY]",
    "",
    "━━━━━━━━━━",
    "",
]

if anchor_ai:
    insert_before(anchor_ai, NEW_STEPS)
    print("✅ Step 7 & 8 插入完成")
else:
    # Fallback：找第六步段落後插入
    for p in doc.paragraphs:
        if "第六步，詢問照片或補充照片" in para_text(p):
            parent = p._p.getparent()
            idx = list(parent).index(p._p)
            for i, line in enumerate(NEW_STEPS):
                parent.insert(idx + 1 + i, make_para(line))
            print("✅ Step 7 & 8 插入完成（fallback 錨點）")
            break

# ── 更新 AI 分析區塊標題，標記為選用 ────────────────────────────────────────
for p in doc.paragraphs:
    txt = para_text(p)
    if "【第七步：AI 照片分析" in txt:
        for run in p.runs:
            if "第七步：AI 照片分析" in run.text:
                run.text = run.text.replace(
                    "【第七步：AI 照片分析 → 觸發服務選項字卡】",
                    "【選用｜真人觸發：AI 照片分析 → 服務選項字卡（[SHOW_SERVICE_OPTIONS]）】"
                )
        print("✅ AI 分析區塊標題更新為選用")
        break

doc.save(DEST)
print(f"\n✅ 儲存完成：{DEST}")
