from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

def add_heading(doc, text, size=16, bold=True, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_body(doc, text, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Title block ──────────────────────────────────────────────
add_heading(doc, "Superior PDR 教育課程推薦影片", size=18, bold=True,
            color=(30, 90, 160), align=WD_ALIGN_PARAGRAPH.CENTER)
add_body(doc, "推薦旁白逐字稿　｜　招生受眾版", size=11, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_body(doc, "說話人：新竹東泰高中 黃士騰 董事", size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_hr(doc)

# ── Meta info ────────────────────────────────────────────────
doc.add_paragraph()
add_body(doc, "📋 文稿資訊", size=11, bold=True, space_after=2)

table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
meta = [
    ("影片長度", "約 38–42 秒"),
    ("說話人", "黃士騰 董事（新竹東泰高中）"),
    ("目標受眾", "有意報名 PDR 課程之潛在學員"),
    ("語言", "中文（含一句英文結尾）"),
]
for i, (k, v) in enumerate(meta):
    table.cell(i, 0).text = k
    table.cell(i, 1).text = v
    for cell in table.row_cells(i):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)

doc.add_paragraph()
add_hr(doc)

# ── Script body ───────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "逐字稿", size=14, bold=True, color=(30, 90, 160))
doc.add_paragraph()

lines = [
    "大家好，我是新竹東泰高中黃士騰董事。",
    "",
    "小馬對技術和教育的熱情，是大家有目共睹的。",
    "",
    "他創立卓越凹痕修復中心，也是台灣第一個把 PDR 技術帶進校園的人。",
    "從學生到專業技師，他影響的不只是技術，而是一整個世代的選擇。",
    "",
    "如果你正在思考未來的方向，",
    "Superior PDR 是一扇真正通往國際舞台的門。",
    "",
    "小馬，keep doing what you do。",
    "台灣需要你，我們也需要更多像你一樣的人。",
]

# Blockquote-style box
for line in lines:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(line)
    run.font.size = Pt(13)
    if "keep doing" in line:
        run.italic = True

doc.add_paragraph()
add_hr(doc)

# ── Timing guide ─────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "時間軸參考（剪輯用）", size=13, bold=True, color=(30, 90, 160))
doc.add_paragraph()

timing = [
    ("00:00 – 00:05", "大家好，我是新竹東泰高中黃士騰董事。"),
    ("00:05 – 00:10", "小馬對技術和教育的熱情，是大家有目共睹的。"),
    ("00:10 – 00:22", "他創立卓越凹痕修復中心……一整個世代的選擇。"),
    ("00:22 – 00:32", "如果你正在思考未來的方向……通往國際舞台的門。"),
    ("00:32 – 00:42", "小馬，keep doing what you do。台灣需要你……"),
]

t2 = doc.add_table(rows=len(timing) + 1, cols=2)
t2.style = 'Table Grid'
header_cells = t2.rows[0].cells
header_cells[0].text = "時間"
header_cells[1].text = "內容"
for cell in header_cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(11)

for i, (ts, content) in enumerate(timing, 1):
    t2.cell(i, 0).text = ts
    t2.cell(i, 1).text = content
    for cell in t2.row_cells(i):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)

doc.add_paragraph()
add_hr(doc)

# ── Notes ────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "拍攝建議", size=12, bold=True)
notes = [
    "建議黃董事長以輕鬆自然的語氣口述，避免唸稿感",
    "背景可搭配東泰高中校園或 Superior PDR 教學場景 B-roll",
    "最後「keep doing what you do」建議以英文語氣自然帶出",
    "可搭配字幕輸出，讓受眾閱讀重點",
]
for note in notes:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(note)
    run.font.size = Pt(11)

# ── Footer ───────────────────────────────────────────────────
doc.add_paragraph()
add_body(doc, "© 2026 Superior PDR | 卓越凹痕修復中心", size=9,
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)

out = r"C:\Users\User\Downloads\Myke_Agent\SUPERIOR教育部門推薦影片\東泰高中_推薦旁白_黃士騰董事版.docx"
doc.save(out)
print(f"Saved: {out}")
